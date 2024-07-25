from selenium import webdriver
import os
import requests
from PIL import Image
import time
import pandas as pd
from docx import Document
import glob
from os import path


path = "C:\Program Files (x86)\chromedriver.exe"
driver = webdriver.Chrome(path)
url = "https://www.google.com"
driver.get(url)


class GoogleImageScraper:
    def __init__(
        self,
        webdriver_path,
        image_path,
        search_key="cat",
        number_of_images=1,
        headless=False,
        min_resolution=(0, 0),
        max_resolution=(1920, 1080),
    ):
        if type(number_of_images) != int:
            print("[ERROR] The number of images must be integer.")
            return
        if not os.path.exists(image_path):
            print("[INFO] Image path not found. Created a new folder.")
            os.makedirs(image_path)
        self.driver = driver
        self.search_key = search_key
        self.number_of_images = number_of_images
        self.webdriver_path = webdriver_path
        self.image_path = image_path
        self.url = (
            "https://www.google.com/search?q=%s&source=lnms&tbm=isch&sa=X&ved=2ahUKEwie44_AnqLpAhUhBWMBHUFGD90Q_AUoAXoECBUQAw&biw=1920&bih=947"
            % (search_key)
        )
        self.headless = headless
        self.min_resolution = min_resolution
        self.max_resolution = max_resolution
        self.saved_extension = "jpg"
        self.valid_extensions = ["jpg", "png", "jpeg"]

    def find_image_urls(self):
        print("[INFO] Scraping for image link... Please wait.")
        image_urls = []
        image_alts = []
        image_titles = []
        image_hrefs = []
        count = 0
        missed_count = 0
        self.driver.get(self.url)
        time.sleep(5)
        for indx in range(1, self.number_of_images + 1):
            try:
                imgurl = self.driver.find_element_by_xpath(
                    '//*[@id="islrg"]/div[1]/div[%s]/a[1]/div[1]/img' % (str(indx))
                )
                imgtitlevehref = self.driver.find_element_by_xpath(
                    '//*[@id="islrg"]/div[1]/div[%s]/a[2]' % (str(indx))
                )
                imgurl.click()
                missed_count = 0
            except Exception:
                missed_count = missed_count + 1
                if missed_count > 10:
                    print("[INFO] No more photos.")
                    break
                else:
                    continue

            try:
                time.sleep(1)
                class_names = ["n3VNCb"]
                images = [
                    self.driver.find_elements_by_class_name(class_name)
                    for class_name in class_names
                    if len(self.driver.find_elements_by_class_name(class_name)) != 0
                ][0]
                for image in images:
                    if image.get_attribute("src")[:4].lower() in ["http"]:
                        print("[BİLGİ] %d. %s" % (count, image.get_attribute("src")))
                        image_urls.append(image.get_attribute("src"))
                        image_alts.append(image.get_attribute("alt"))
                        image_titles.append(imgtitlevehref.get_attribute("title"))
                        image_hrefs.append(imgtitlevehref.get_attribute("href"))
                        count += 1
                        break
            except Exception:
                print("[INFO] Unable to receive connection.")

            try:
                if count % 3 == 0:
                    self.driver.execute_script(
                        "window.scrollTo(0, " + str(indx * 60) + ");"
                    )
                element = self.driver.find_element_by_class_name("mye4qd")
                element.click()
                print("[INFO] Uploading more photos")
                time.sleep(5)
            except Exception:
                time.sleep(1)
        self.driver.quit()
        print("[INFO] Google search is over")

        df_data = pd.DataFrame(
            {
                "Src": image_urls,
                "Alt": image_alts,
                "Href": image_hrefs,
                "Title": image_titles,
            }
        )
        df_data
        df_data.to_excel("Search Results.xlsx", index=False)
        return image_urls

    def save_images(self, image_urls):
        print("[INFO] Saving Picture... Please wait...")
        for indx, image_url in enumerate(image_urls):
            try:
                filename = "%s%s.%s" % (
                    self.search_key,
                    str(indx),
                    self.saved_extension,
                )
                image_path = os.path.join(self.image_path, filename)
                print("[INFO] %d.Image saved at: %s" % (indx, image_path))
                image = requests.get(image_url)
                if image.status_code == 200:
                    with open(image_path, "wb") as f:
                        f.write(image.content)
                        f.close()
                        image_from_web = Image.open(image_path)
                        image_resolution = image_from_web.size
                        if image_resolution != None:
                            if (
                                image_resolution[0] < self.min_resolution[0]
                                or image_resolution[1] < self.min_resolution[1]
                                or image_resolution[0] > self.max_resolution[0]
                                or image_resolution[1] > self.max_resolution[1]
                            ):
                                image_from_web.close()
                                os.remove(image_path)
                        image_from_web.close()
            except Exception as e:
                print("[ERROR] Failed to download", e)
                pass
        print(
            "[INFO] Download Completed. Please note that some photos were not downloaded because they were not in the correct format. (e.g. jpg, jpeg, png)"
        )


webdriver_path = os.path.normpath(os.getcwd() + "\\webdriver\\chromedriver.exe")
image_path = os.path.normpath(os.getcwd() + "\\images")

search_keys = ["sustainable fashion product"]

number_of_images = 5
headless = False
min_resolution = (0, 0)
max_resolution = (9999, 9999)

for search_key in search_keys:
    image_scrapper = GoogleImageScraper(
        webdriver_path,
        image_path,
        search_key,
        number_of_images,
        headless,
        min_resolution,
        max_resolution,
    )
    image_urls = image_scrapper.find_image_urls()
    image_scrapper.save_images(image_urls)

del image_scrapper

fotograf = glob.glob("./images/*.jpg")
document = Document()

for file in fotograf:
    p = document.add_paragraph(file)
    r = p.add_run()
    r.add_picture(file)

document.save("Images.docx")
